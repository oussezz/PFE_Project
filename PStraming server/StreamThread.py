import cv2
import imutils
import io
from imutils.video import VideoStream
import threading
import zmq
import cv2
import datetime
import base64
import os



import requests
import json
import base64

socket=[]

endpoint_url = 'http://localhost:4040/predict'




class StreamThread:
    
    def create_report(self):
        print("creating pdf : "+str(self.add))
        #create a notification
        import locale
        from docx import Document
        from docx.shared import Inches
        import random
        
        import docx2pdf
        import json
        from datetime import datetime
        # Set the system's locale to Arabic
        locale.setlocale(locale.LC_ALL, 'ar')
        # Replace the variables with the desired values
        now = datetime.now()
        date_string = now.strftime('%d %B %Y')
        time_string = now.strftime('%I:%M:%S %p')
        lat = random.uniform(36, 37)
        lon = random.uniform(2, 3)
        #create file if doesnt exist
        parent_dir = os.path.abspath(os.path.join(os.getcwd(), os.pardir))
        parent_dir=os.path.join(parent_dir,'FDPlatform/mediafiles/reports/'+str(self.add)+'/')
        #create the directory if it does not exist
        if not os.path.exists(parent_dir):
            os.makedirs(parent_dir)
        # Save the modified Word document as a PDF
        
    # Set the system's locale to Arabic
        locale.setlocale(locale.LC_ALL, 'ar')
        # Open the Word template
        doc = Document('template.docx')
        # Set the system's locale to Arabic
        locale.setlocale(locale.LC_ALL, 'ar')
        pdf_name =str(self.add)+'___'+ now.strftime('%Y-%m-%d_%H-%M-%S') + '.pdf'
        if not os.path.exists("C:/Users/Ezziane/FireDetectionTemp/notif/notif.json"):
            json_file = open("C:/Users/Ezziane/FireDetectionTemp/notif/notif.json", 'w')
            json_file.write('{"notifications":[{"id": "2", "date": "2023-04-30 10:07:57.927488", "type": "pwd", "description": "Your password has been updated successfully", "status": "read", "path": ""}]}')
            json_file.close()
        
        with open("C:/Users/Ezziane/FireDetectionTemp/notif/notif.json", 'r') as f:
        # write dictionary to the file
            data = json.load(f)
        notifications=data['notifications']
        last_id = notifications[-1]['id']
        new_id = int(last_id) + 1

        data['notifications'].append({'id':str(new_id),'date':str(datetime.now()),"type": "report",'description':'Fire detected in '+str(self.add),'status':'unread','path':'/media/reports/'+str(self.add)+'/'+pdf_name})
        with open("C:/Users/Ezziane/FireDetectionTemp/notif/notif.json", 'w') as f:
            json.dump(data, f)    
        
        
        

        def replace_text(doc, old_text, new_text):
            """
            Replace all occurrences of old_text with new_text in the given document.
            """
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text(cell, old_text, new_text)

        

        # Set the system's locale to Arabic
        locale.setlocale(locale.LC_ALL, 'ar')
        replace_text(doc, '[DATE]', date_string)
        replace_text(doc, '[TIME]', time_string)
        replace_text(doc, '[LAT]', str(lat))
        replace_text(doc, '[LON]', str(lon))
        doc.save('fire_reportDONE.docx')
        
        docx2pdf.convert('fire_reportDONE.docx',parent_dir+ pdf_name)
        
    
    
    
    def stream(self,stop_event,add):
            global socket
            i=0
            # port = int(add.split('.')[-1])+2000
            port = 2000
            
            if not socket:
                context = zmq.Context()
                socket = context.socket(zmq.PUB)
                socket.bind("tcp://*:"+str(port))
            # video_stream = cv2.VideoCapture('http://'+str(add)+':8080/video')
            video_stream = cv2.VideoCapture('0001-2102.mp4')
            print(f'streaming...http://'+str(add)+':8080/video')
            elapsed_time = datetime.timedelta(seconds=0)
            timer_15_min = datetime.timedelta(minutes=10)
            parent_dir = os.path.abspath(os.path.join(os.getcwd(), os.pardir))
            parent_dir=os.path.join(parent_dir,'FDPlatform/mediafiles/history/cameras/'+str(add))
            #create the directory if it does not exist
            if not os.path.exists(parent_dir):
                os.makedirs(parent_dir)
                
            out = None
            fourcc = cv2.VideoWriter_fourcc(*'mp4v')
            fps = 30.0
            while not stop_event.is_set():
                rep,frame = video_stream.read()
                if frame is None:
                    continue
                if rep:
                    # If video writer object does not exist or 15 minutes have passed, create a new video file
                    if out is None or elapsed_time >= timer_15_min:
                        # Release previous video writer object, if it exists
                        if out is not None:
                            out.release()

                        # Create new video writer object with timestamped name in parent directory
                        timestamp = datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
                        file_path = os.path.join(parent_dir, '{}.mp4'.format(timestamp))
                        out = cv2.VideoWriter(file_path, fourcc, fps, (640, 480))

                        # Reset elapsed time
                        elapsed_time = datetime.timedelta(seconds=0)

                    # Write frame to video writer object
                    out.write(frame)

                    # Increment elapsed time
                    elapsed_time += datetime.timedelta(seconds=1)
                
                now = datetime.datetime.now()

               

                # Check if the seconds are even
                def send_request(img):
                    _, image_bytes = cv2.imencode('.jpg', img)
                    image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                    
                    # Set the data for the request
                    data = {'image': image_base64,  'ip': str(add)}
                    
                    # Send the POST request
                    requests.post(endpoint_url, json=data,stream=True)
                    
                    
                i+=1
                if i % 60 == 0:
                    threading.Thread(target=send_request,args=(frame,)).start()
                    i=0
                    
                #open the json file
                try:
                    if os.path.isfile("C:/Users/Ezziane/FireDetectionTemp/"+str(add)+".json"):
                        with open("C:/Users/Ezziane/FireDetectionTemp/"+str(add)+".json", 'r') as f:
                            data = json.load(f)
                        class_label = data['class']
                        probability = data['probabilities']
                        if class_label == 'Fire':
                            #check if a file exist
                            if not os.path.isfile("C:/Users/Ezziane/FireDetectionTemp/should_create_a_report_"+str(add)+".json"):
                                #create a report
                                threading.Thread(target=self.create_report,).start()
                                #create a file
                                with open("C:/Users/Ezziane/FireDetectionTemp/should_create_a_report_"+str(add)+".json", 'w') as f:
                                    json.dump({'should_create_a_report':True}, f)
                                    

                            
                            
                            
                            
                            
                        # Write the class label and probability on the image
                        cv2.putText(frame, f"Class: {class_label}", (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 255, 0), 1)
                        cv2.putText(frame, f"Prob: {probability:.2f}", (10, 50), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 255, 0), 1)

                except:
                    pass
                # Get the current date and time
                now = datetime.datetime.now()
                date_str = now.strftime("%Y-%m-%d")
                time_str = now.strftime("%H:%M:%S")

                # Set the font and font scale
                font = cv2.FONT_HERSHEY_SIMPLEX
                font_scale = 0.8

                # Set the text and text position
                text = f"{date_str} {time_str}"
                text_pos = (frame.shape[1] - 300, 20)

                # Set the text color, thickness, and line type
                text_color = (255, 255, 255)
                text_thickness = 2
                line_type = cv2.LINE_AA

                # Draw the text on the frame
                cv2.putText(frame, text, text_pos, font, font_scale, text_color, text_thickness, line_type)


                # frame = cv2.resize(frame,(720,480))
                frame= cv2.imencode(".jpg", frame)[1].tobytes()
                label = f'{add}'.encode('utf-8')
                message= label + b' ' + frame   
                # msg=[str(add),frame]
                socket.send(message)
            video_stream.release()
            cv2.destroyAllWindows()
    def __init__(self,add,inference_server):
        self.add=str(add)
        self.stop_event = threading.Event()
        self.thread=threading.Thread(target=self.stream,args=(self.stop_event,add,))
        self.thread.start()
        self.inference_server=inference_server
        
        
        